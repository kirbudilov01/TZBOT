import logging
import sqlite3
import os
import asyncio
from aiogram import Bot, Dispatcher, types
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters.state import State, StatesGroup
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from docx import Document
import openai

# API ключи из переменных окружения
API_TOKEN = os.getenv("API_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Логирование
logging.basicConfig(level=logging.INFO)

# Инициализация бота и диспетчера
bot = Bot(token=API_TOKEN)
dp = Dispatcher(bot, storage=MemoryStorage())

# Подключение к базе данных SQLite
conn = sqlite3.connect("tz_database.db")
cursor = conn.cursor()
cursor.execute('''CREATE TABLE IF NOT EXISTS requests (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    user_id INTEGER,
    business_goal TEXT,
    key_features TEXT,
    integrations TEXT,
    target_audience TEXT,
    monetization TEXT,
    security TEXT,
    technologies TEXT,
    additional_questions TEXT,
    generated_tz TEXT
)''')
conn.commit()

# Определяем состояния
class Form(StatesGroup):
    business_goal = State()
    key_features = State()
    integrations = State()
    target_audience = State()
    monetization = State()
    security = State()
    technologies = State()
    additional = State()

# Команда /start
@dp.message_handler(commands=['start'])
async def start(message: types.Message):
    await message.answer("Привет! Я помогу составить детальное техническое задание. Давайте начнем!\n\nКакова основная цель вашего продукта? (Например: автоматизация продаж, чат-бот для поддержки, маркетинговый инструмент и т.д.)")
    await Form.business_goal.set()

# Вопросы по продукту
@dp.message_handler(state=Form.business_goal)
async def process_business_goal(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['business_goal'] = message.text
    await message.answer("Какие основные функции должны быть в приложении/боте? Опишите максимально подробно.")
    await Form.key_features.set()

@dp.message_handler(state=Form.key_features)
async def process_key_features(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['key_features'] = message.text
    await message.answer("Какие сервисы и API нужно интегрировать? (например, CRM, платежные системы, OpenAI и т.д.)")
    await Form.integrations.set()

@dp.message_handler(state=Form.integrations)
async def process_integrations(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['integrations'] = message.text
    await message.answer("Кто ваша основная целевая аудитория? (Малый бизнес, B2B, B2C, маркетологи и т.д.)")
    await Form.target_audience.set()

@dp.message_handler(state=Form.target_audience)
async def process_target_audience(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['target_audience'] = message.text
    await message.answer("Как планируется монетизация? (Подписка, разовая покупка, реклама и т.д.)")
    await Form.monetization.set()

@dp.message_handler(state=Form.monetization)
async def process_monetization(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['monetization'] = message.text
    await generate_tz(message, state)

# Генерация развернутого ТЗ через GPT-4
async def generate_tz(message, state):
    async with state.proxy() as data:
        prompt = f"""
        Составь детальное техническое задание на основе следующих данных:
        Цель продукта: {data['business_goal']}
        Основные функции: {data['key_features']}
        Интеграции: {data['integrations']}
        Целевая аудитория: {data['target_audience']}
        Монетизация: {data['monetization']}
        
        Используй формат:
        1. Введение
        2. Описание продукта
        3. Функционал
        4. Технологический стек
        5. Монетизация
        6. Безопасность
        
        Сделай текст развернутым, как в примере документа.
        """
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}]
        )
        generated_text = response["choices"][0]["message"]["content"]
        
        document = Document()
        document.add_heading("Техническое задание", level=1)
        document.add_paragraph(generated_text)
        
        file_path = f"tz_{message.from_user.id}.docx"
        document.save(file_path)
        
        with open(file_path, "rb") as file:
            await message.answer("Спасибо! Вот ваше детальное техническое задание:")
            await bot.send_document(message.chat.id, file)
        
        os.remove(file_path)
    await state.finish()

# Запуск бота
async def main():
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())
